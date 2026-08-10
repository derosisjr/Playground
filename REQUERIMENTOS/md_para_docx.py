# -*- coding: utf-8 -*-
"""Converte as peças em Markdown desta pasta para .docx protocolável.

Existe porque as peças nascem em Markdown (versionável, diffável, fácil de
revisar) mas são protocoladas em Word. Não há pandoc nesta máquina; o
`python-docx` já está instalado e dá conta do subconjunto que as peças usam:
títulos, parágrafos, **negrito**, listas numeradas e tabelas.

Deliberadamente NÃO converte a "Nota de apuração": ela é de uso interno do
gabinete — traz verificações pendentes e o que não foi comprovado — e não pode
sair junto no protocolo. O corte é feito no marcador `### Nota de apuração`.

Uso:
  python REQUERIMENTOS/md_para_docx.py piso-moeda/Requerimento_Piso_Moeda_Pregao_96-2025.md
  python REQUERIMENTOS/md_para_docx.py --todas
"""
import argparse
import glob
import os
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

DIR = os.path.dirname(os.path.abspath(__file__))
CORTE = "### Nota de apuração"


def _desescapar(s):
    """`\\_` e `\\*` do Markdown voltam ao caractere puro.

    Aplicado em TODOS os ramos: a linha de preenchimento do número do
    requerimento ("Nº \\_\\_\\_\\_/2026") está inteira em negrito, e desescapar só
    no ramo de texto simples deixava as barras invertidas no documento final.
    """
    return s.replace("\\_", "_").replace("\\*", "*")


def _inline(par, texto):
    """Escreve o texto no parágrafo, resolvendo **negrito** e `código`."""
    for pedaco in re.split(r"(\*\*.+?\*\*|`.+?`)", texto):
        if not pedaco:
            continue
        if pedaco.startswith("**") and pedaco.endswith("**"):
            par.add_run(_desescapar(pedaco[2:-2])).bold = True
        elif pedaco.startswith("`") and pedaco.endswith("`"):
            r = par.add_run(_desescapar(pedaco[1:-1]))
            r.font.name = "Consolas"
        else:
            par.add_run(_desescapar(pedaco))


def _tabela(doc, linhas):
    celulas = [[c.strip() for c in l.strip().strip("|").split("|")] for l in linhas]
    celulas = [c for c in celulas if not all(set(x) <= set("-: ") for x in c)]
    if not celulas:
        return
    t = doc.add_table(rows=0, cols=len(celulas[0]))
    t.style = "Table Grid"
    for i, linha in enumerate(celulas):
        cells = t.add_row().cells
        for j, txt in enumerate(linha[:len(cells)]):
            cells[j].text = ""
            _inline(cells[j].paragraphs[0], txt)
            if i == 0:
                for r in cells[j].paragraphs[0].runs:
                    r.bold = True
    doc.add_paragraph()


def converter(caminho_md, saida=None):
    with open(caminho_md, encoding="utf-8") as f:
        texto = f.read()
    if CORTE in texto:
        texto = texto.split(CORTE)[0]          # nota interna nunca vai ao protocolo
    texto = texto.rstrip().rstrip("-").rstrip()

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    linhas = texto.split("\n")
    i, buffer_tabela = 0, []
    while i < len(linhas):
        l = linhas[i]
        if l.lstrip().startswith("|"):
            buffer_tabela.append(l)
            i += 1
            continue
        if buffer_tabela:
            _tabela(doc, buffer_tabela)
            buffer_tabela = []
        s = l.strip()
        if not s or s == "---":
            i += 1
            continue
        if s.startswith("## "):
            doc.add_heading(s[3:].strip(), level=1)
        elif s.startswith("# "):
            doc.add_heading(s[2:].strip(), level=0)
        elif re.match(r"^\d+\.\s", s):
            corrido = [re.sub(r"^\d+\.\s", "", s)]
            j = i + 1
            while j < len(linhas) and linhas[j].startswith(("   ", "	")) and linhas[j].strip():
                corrido.append(linhas[j].strip()); j += 1
            p = doc.add_paragraph(style="List Number")
            _inline(p, " ".join(corrido))
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            i = j
            continue
        elif s.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _inline(p, s[2:])
        else:
            # Markdown quebra a linha por largura de coluna; no Word isso tem de
            # voltar a ser UM parágrafo, senão cada linha do fonte vira um
            # parágrafo solto e a justificação não funciona
            corrido = [s]
            j = i + 1
            while j < len(linhas):
                prox = linhas[j].strip()
                if (not prox or prox == "---" or prox.startswith(("#", "- ", "|"))
                        or re.match(r"^\d+\.\s", prox)):
                    break
                corrido.append(prox)
                j += 1
            p = doc.add_paragraph()
            _inline(p, " ".join(corrido))
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            i = j
            continue
        i += 1
    if buffer_tabela:
        _tabela(doc, buffer_tabela)

    saida = saida or os.path.splitext(caminho_md)[0] + ".docx"
    doc.save(saida)
    return saida


def main():
    ap = argparse.ArgumentParser(description="Markdown → .docx para protocolo")
    ap.add_argument("arquivo", nargs="?", help="caminho relativo a REQUERIMENTOS/")
    ap.add_argument("--todas", action="store_true", help="converte as peças com nota interna")
    args = ap.parse_args()
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")

    if args.todas:
        alvos = [p for p in glob.glob(os.path.join(DIR, "**", "*.md"), recursive=True)
                 if CORTE in open(p, encoding="utf-8").read()]
    elif args.arquivo:
        alvos = [os.path.join(DIR, args.arquivo)]
    else:
        raise SystemExit("informe um arquivo ou use --todas")

    for p in alvos:
        print(f"{os.path.relpath(p, DIR)} → {os.path.basename(converter(p))}")


if __name__ == "__main__":
    main()
