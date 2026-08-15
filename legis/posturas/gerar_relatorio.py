# -*- coding: utf-8 -*-
"""
Gera relatorio_posturas.docx — relatório da Comissão Especial sobre o PLC do
novo Código de Posturas de Santos (PA 59314/2025-78) × Lei nº 3.531/1968.
Parte A: descrição neutra. Parte B: pontos de atenção (perguntas e emendas).
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
GOLD = RGBColor(0xC9, 0xA2, 0x4B)
BASE = Path(__file__).parent
SAIDA = BASE / "relatorio_posturas.docx"

doc = Document()

# corrige w:zoom sem atributo obrigatório w:percent (template padrão do python-docx)
_zoom = doc.settings.element.find(qn("w:zoom"))
if _zoom is not None and _zoom.get(qn("w:percent")) is None:
    _zoom.set(qn("w:percent"), "100")

# ---- página A4, margens 2,5 cm ----
sec = doc.sections[0]
sec.page_width, sec.page_height = Mm(210), Mm(297)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Cm(2.5))

# ---- estilos ----
normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for nome, tam, cor in (("Heading 1", 16, NAVY), ("Heading 2", 13, NAVY),
                       ("Heading 3", 11.5, GOLD)):
    st = doc.styles[nome]
    st.font.name = "Arial"
    st.font.size = Pt(tam)
    st.font.bold = True
    st.font.color.rgb = cor

# rodapé com nº de página
foot_p = sec.footer.paragraphs[0]
foot_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = foot_p.add_run()
for tag, attrs, text in (("w:fldChar", {"w:fldCharType": "begin"}, None),
                         ("w:instrText", {"xml:space": "preserve"}, " PAGE "),
                         ("w:fldChar", {"w:fldCharType": "end"}, None)):
    el = OxmlElement(tag)
    for k, v in attrs.items():
        el.set(qn(k), v)
    if text:
        el.text = text
    run._r.append(el)
run.font.size = Pt(9)


def p(texto, bold=False, italic=False, size=None, align=None, color=None):
    par = doc.add_paragraph()
    r = par.add_run(texto)
    r.bold, r.italic = bold, italic
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    if align:
        par.alignment = align
    return par


def bullet(texto, nivel=0):
    par = doc.add_paragraph(style="List Bullet" if nivel == 0 else "List Bullet 2")
    par.add_run(texto)
    return par


def rotulado(rotulo, texto):
    par = doc.add_paragraph()
    r = par.add_run(rotulo + " ")
    r.bold = True
    r.font.color.rgb = NAVY
    par.add_run(texto)
    return par


def tabela(cabecalho, linhas, larguras_cm):
    t = doc.add_table(rows=1, cols=len(cabecalho))
    t.style = "Table Grid"
    for i, (h, w) in enumerate(zip(cabecalho, larguras_cm)):
        cel = t.rows[0].cells[i]
        cel.width = Cm(w)
        par = cel.paragraphs[0]
        r = par.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "1B2A4A")
        cel._tc.get_or_add_tcPr().append(shd)
    for linha in linhas:
        cells = t.add_row().cells
        for i, (v, w) in enumerate(zip(linha, larguras_cm)):
            cells[i].width = Cm(w)
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(10)
    return t


# ================================================================= CAPA
for _ in range(6):
    doc.add_paragraph()
p("RELATÓRIO DE ANÁLISE", bold=True, size=24, align=WD_ALIGN_PARAGRAPH.CENTER, color=NAVY)
p("Projeto de Lei Complementar — Novo Código de Posturas do Município de Santos",
  bold=True, size=15, align=WD_ALIGN_PARAGRAPH.CENTER)
p("PA 59314/2025-78 — comparativo com a Lei nº 3.531, de 16 de abril de 1968 (texto consolidado)",
  size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GOLD)
for _ in range(3):
    doc.add_paragraph()
p("Comissão Especial de Modernização do Código de Posturas", bold=True, size=12,
  align=WD_ALIGN_PARAGRAPH.CENTER)
p("Câmara Municipal de Santos — Gabinete da Presidência da Comissão", size=11,
  align=WD_ALIGN_PARAGRAPH.CENTER)
p("Julho de 2026", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(4):
    doc.add_paragraph()
p("Documento de apoio parlamentar elaborado por leitura comparada dos dois textos. "
  "Não constitui parecer jurídico. Itens não confirmados em fonte primária estão "
  "marcados [Verificar].", italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ======================================================= SUMÁRIO EXECUTIVO
doc.add_heading("Sumário executivo", level=1)
bullet("O PLC substitui integralmente o Código de Posturas de 1968 — cerca de 640 artigos "
       "consolidados ao longo de 57 anos — por um texto novo de 335 artigos, revogando "
       "nominalmente 23 normas (art. 334).")
bullet("A proposta tem natureza tripla: (i) desregulamentação e enxugamento (corte de ~48% dos "
       "dispositivos, remoção de matérias de competência federal/estadual); (ii) modernização da "
       "fiscalização e do sistema sancionatório (multas em UFESP com dosimetria, notificação "
       "eletrônica, termo de compromisso); (iii) incorporação de temas novos (drenagem urbana, "
       "zoonoses, parques, concessionárias e cabeamento).")
bullet("Destino dos 640 artigos do código vigente: 202 alterados no mérito, 145 condensados, "
       "142 remetidos a outras normas, 82 revogados sem correspondente, 49 já revogados por leis "
       "anteriores e 20 mantidos em substância (detalhes na planilha comparativo_posturas.xlsx).")
bullet("A direção geral é compatível com a agenda de desburocratização. Merecem exame da "
       "Comissão: a forte delegação de matérias relevantes a decreto, o endurecimento "
       "sancionatório em pontos específicos, a ausência de disposições transitórias, lacunas "
       "materiais (ex.: inflamáveis e postos de combustível) e a cláusula revogatória incompleta.")
bullet("Recomenda-se oficiar o Executivo com as perguntas consolidadas na seção 4 e preparar "
       "emendas pontuais indicadas na Parte B.")

# ============================================================ 1. MÉTODO
doc.add_heading("1. Objeto e método", level=1)
p("Foram comparados o texto consolidado da Lei nº 3.531/1968 (versão LeisMunicipais, com "
  "alterações até 06/01/2025) e o PLC encaminhado pelo Executivo (PA 59314/2025-78). Cada um "
  "dos 640 artigos do código vigente foi classificado quanto ao seu destino no PLC (mantido, "
  "alterado, condensado, remetido a outra norma, revogado sem correspondente ou já revogado "
  "anteriormente), e os dispositivos novos do PLC foram catalogados à parte. O detalhamento "
  "integral — inclusive o mapa artigo a artigo — está na planilha comparativo_posturas.xlsx, "
  "anexa a este relatório. Artigos já revogados por leis anteriores (ex.: cemitérios, "
  "LC 712/2011; ambulantes, LC 1189/2023) não foram computados como cortes do PLC.")

# ============================================================ 2. PARTE A
doc.add_heading("2. Parte A — O que muda (descrição objetiva)", level=1)

doc.add_heading("2.1 Nova moldura de princípios e direitos (PLC, arts. 1º–10)", level=2)
p("O PLC abre com objetivos e princípios ausentes do código atual: higiene e salubridade, "
  "ordem e sossego, proteção do meio ambiente, acessibilidade, desburocratização e "
  "simplificação (art. 2º, XI) e livre iniciativa (art. 3º, XI), ao lado de função "
  "socioambiental da propriedade, direito à cidade sustentável e gestão democrática (art. 3º). "
  "Positiva direitos dos munícipes (devido processo, contraditório, informação — art. 10) e o "
  "dever de colaborar com a fiscalização, com multa de 60 a 600 UFESP por impedir o acesso dos "
  "agentes (art. 5º, §3º).")

doc.add_heading("2.2 Higiene pública condensada (PLC, arts. 11–104)", level=2)
p("O Título II atual (181 artigos) é reduzido a 94. As minúcias de 1968 — balcões de mármore, "
  "açucareiros, telas, iluminação de açougues — dão lugar a deveres gerais de higiene com "
  "remissão à legislação sanitária. Permanecem quase literais as definições de alimento "
  "impróprio, adulterado e fraudado (art. 34). São novos: o capítulo de acondicionamento de "
  "resíduos, coleta seletiva e cata-treco (arts. 52–57), o controle de zoonoses (arts. 58–64) e "
  "o capítulo ampliado de controle da poluição — ar (escala Ringelmann), odores, ruído por NBR "
  "10151/10152, solo e águas — com multas de 200 a 2.000 UFESP (arts. 69–90). Na limpeza de "
  "terrenos, a multa fixa atual de R$ 15.067,25 é substituída por multa diária de 380 a 432 "
  "UFESP até a conclusão dos serviços (art. 101), mantida a execução pela Prefeitura com taxa "
  "de administração de 100%.")

doc.add_heading("2.3 Título novo: drenagem urbana (PLC, arts. 105–127)", level=2)
p("Sem paralelo no código atual. Impõe: sistema interno de captação de águas pluviais com "
  "soluções de retenção (art. 122); vedação de lançamento de esgoto na rede pluvial e de águas "
  "pluviais na rede de esgoto (art. 123); caixas de gordura com limpeza periódica certificada "
  "(art. 115); deveres de concessionárias (arts. 117–119) e de construtoras — controle de "
  "erosão, recomposição e corresponsabilidade por impactos por 5 anos após a entrega da obra "
  "(arts. 120–121). Multas escalonadas chegam a 20.000 UFESP para abatimento de via superior a "
  "50 m² (art. 127, V).")

doc.add_heading("2.4 Bem-estar público (PLC, arts. 128–220)", level=2)
p("Sossego: os limites deixam de constar em lei e passam às NBR 10151/10152; o rol de "
  "tolerâncias é tipificado (ensaios de carnaval, artistas de rua, cultos, obras das 8h às 18h "
  "em dias úteis — art. 138); mantida a proibição de fogos de estampido, agora vedada também a "
  "instalação de estabelecimentos que os produzam ou comercializem (art. 139). Praias: prática "
  "esportiva, barracas e uso geral passam a depender de decreto (arts. 147, 149, 154); "
  "proibidos equipamentos de som (art. 148, IV) e a captura do corrupto — Callichirus (art. "
  "153). Novo capítulo de parques municipais (arts. 156–162). Calçadas: responsabilidade do "
  "proprietário mantida; a largura da faixa livre passa a ser definida em regulamento (art. "
  "167, I). Publicidade: de ~20 artigos para 5; caem o limite de 25% da fachada e os prazos de "
  "licença; seguem vedadas colagem em postes e panfletagem em residências, com multa dirigida "
  "ao anunciante (arts. 179–183). Animais: charretes e carroças proibidas em todo o Município "
  "(art. 210, III); circulação de cães na praia em área e horário definidos pelo Executivo "
  "(art. 206); câmeras em banho e tosa mantidas, com guarda de imagens de 30 dias (art. 212).")

doc.add_heading("2.5 Localização e funcionamento (PLC, arts. 221–279)", level=2)
p("Licenciamento com consulta prévia de viabilidade, prazo de 15 dias mantido, exigência de "
  "AVCB e Licença Ambiental para atividades de alto risco e responsabilização do proprietário "
  "do imóvel por uso divergente do cadastro (art. 221, §4º). A emissão do alvará passa a ser "
  "disponibilizada no site da Prefeitura (art. 224), mantida, porém, a obrigação de conservá-lo "
  "afixado em local visível à fiscalização (art. 224, §1º) — ou seja, muda o canal de emissão, "
  "não há regime de alvará dematerializado. O horário de "
  "funcionamento é desregulamentado: a tabela minuciosa por ramo (arts. 435–454 atuais) vira "
  "regra geral de respeito ao sossego (art. 229). Comércio ambulante, feiras e bancas ficam em "
  "legislação específica, com pilares gerais de conduta (arts. 233–236). Diversões públicas: "
  "licença especial para lotação acima de 300 pessoas, Certificado de Conformidade Técnica com "
  "ARTs, e exigências de enfermagem, ambulância e brigada para eventos acima de 1.500 pessoas "
  "(arts. 237–247). Capítulo novo obriga concessionárias à recomposição de logradouros (LC "
  "852/2014), organização de cabeamento aéreo, vedação de novos cabos em postes saturados e "
  "remoção de fiação obsoleta, com multas de 215 a 670 UFESP (arts. 263–279).")

doc.add_heading("2.6 Matérias removidas do Código", level=2)
bullet("Instalações elétricas e mecânicas (arts. 309–426 atuais, 118 artigos: caldeiras, "
       "elevadores, ascensoristas, escadas rolantes) — remissão a ABNT, Corpo de Bombeiros e "
       "Código de Edificações (PLC, arts. 217–220).")
bullet("Inflamáveis, explosivos e postos de combustível (arts. 506–537) — capítulo desaparece, "
       "inclusive a distância mínima de 200 m de escolas e hospitais para novos postos "
       "(art. 536, §3º atual). [Verificar cobertura pela lei de uso do solo e normas "
       "estaduais/federais]")
bullet("Segurança no trabalho (arts. 552–567) e aferição de pesos e medidas (arts. 568–572) — "
       "matérias de competência federal (CLT/NRs; Inmetro/IPEM); a remoção corrige sobreposição.")
bullet("Outras supressões sem correspondente: supermercados (regras próprias), campos "
       "esportivos, locais de culto, tapumes e andaimes, utilização e iluminação de edifícios, "
       "vitrinas e mostruários, balneários, aceitação de instalações, placas de advertência "
       "sobre bebidas alcoólicas e energéticos, Comissão Consultiva do Código (art. 636 atual).")

doc.add_heading("2.7 Fiscalização (PLC, arts. 280–298)", level=2)
p("Tipifica prerrogativas e deveres dos agentes, prevê ciência de notificações por via "
  "pessoal, postal, eletrônica ou edital (art. 289) e cria o termo de compromisso: ajuste de "
  "conduta com força de título executivo extrajudicial que suspende sanções não cautelares "
  "enquanto cumprido, com prazo improrrogável (arts. 293–298).")

doc.add_heading("2.8 Infrações e penalidades (PLC, arts. 299–330)", level=2)
p("O sistema sancionatório é reconstruído: multas em UFESP com dosimetria (multa-base + "
  "atenuantes e agravantes), tratamento diferenciado a pessoa física de baixa renda, MEI, ME e "
  "EPP (art. 300, V), defesa em 30 dias, reincidência caracterizada em 12 meses. São "
  "sistematizados: advertência, apreensão com rito e destinação (retirada em 5 dias; leilão, "
  "doação ou destruição), suspensão de até 60 dias, embargo, interdição, lacração e "
  "emparedamento — com recurso sem efeito suspensivo (arts. 319, §3º e 324, parágrafo único) — "
  "cassação com impedimento de 5 anos alcançando os sócios (art. 327, §2º) e demolição com "
  "execução subsidiária pela Prefeitura, independentemente de ordem judicial, com custo "
  "acrescido de 100% (art. 330).")

doc.add_heading("2.9 Cláusula revogatória (PLC, art. 334)", level=2)
p("Revoga nominalmente 23 normas: a Lei nº 3.531/1968 e 21 leis complementares (365/1999 a "
  "986/2017), além da Lei nº 3.381/2017 [Verificar numeração]. A lista consta da planilha "
  "anexa, aba “Revogações Expressas”.")

# ============================================================ 3. PARTE B
doc.add_heading("3. Parte B — Pontos de atenção da Comissão", level=1)
p("Análise orientada pela defesa do contribuinte, pela desburocratização e pelo controle da "
  "atividade do Executivo. Cada item indica o dispositivo, o impacto prático, a pergunta "
  "sugerida ao Executivo e, quando cabível, a linha de emenda.", italic=True)

doc.add_heading("3.1 Méritos a registrar", level=2)
bullet("Corte de ~48% dos dispositivos e eliminação de minúcias obsoletas de 1968 — alinhado à "
       "agenda de desburocratização e liberdade econômica.")
bullet("Remoção de matérias de competência alheia (segurança do trabalho, metrologia, "
       "rotulagem), reduzindo insegurança jurídica e dupla autuação.")
bullet("Liberdade de horário de funcionamento (art. 229) — fim da tabela por ramo de 1968.")
bullet("Termo de compromisso (arts. 293–298) como alternativa negociada à sanção.")
bullet("Dosimetria com tratamento diferenciado a MEI, ME, EPP e baixa renda (art. 300, V).")
p("Recomendação: registrar apoio expresso a esses eixos no parecer da Comissão.", italic=True)

doc.add_heading("3.2 Delegação excessiva a decreto", level=2)
rotulado("Dispositivos:", "arts. 154 (uso das praias), 167, I (largura da faixa livre das "
         "calçadas), 233 (comércio ambulante/permissionários), 288, VII (prazos das "
         "notificações), 138, §2º (medidas mitigatórias de ruído), 147/149/150 (esportes e "
         "barracas na praia).")
rotulado("Impacto:", "parâmetros hoje fixados em lei — ex.: faixa livre de 1,50 m para mesas e "
         "cadeiras (art. 233, II atual) — passam a regulamento, esvaziando o controle da Câmara "
         "e criando instabilidade para o administrado.")
rotulado("Pergunta ao Executivo:", "quais minutas de decreto acompanharão a sanção e em que "
         "prazo serão editadas?")
rotulado("Emenda sugerida:", "fixar em lei pisos mínimos (faixa livre ≥ 1,50 m; prazo mínimo de "
         "defesa; balizas gerais de uso das praias), delegando ao decreto apenas o detalhamento.")

doc.add_heading("3.3 Carga sancionatória e impacto sobre o contribuinte", level=2)
rotulado("Dispositivos:", "faixas de multa em UFESP em todo o texto; multa diária de 380–432 "
         "UFESP para terrenos (art. 101); até 20.000 UFESP em drenagem (art. 127); 200–2.000 "
         "UFESP em poluição (art. 90).")
rotulado("Impacto:", "com a UFESP de referência [Verificar valor vigente na tramitação], a "
         "multa mínima de higiene (20 UFESP) fica próxima de R$ 740; a multa de terreno sujo, "
         "hoje fixa em R$ 15.067,25, passa a valor equivalente POR DIA de descumprimento — "
         "potencial multiplicador relevante sobre proprietários; sanções de drenagem podem "
         "superar R$ 700 mil.")
rotulado("Pergunta ao Executivo:", "há estudo de impacto e tabela comparativa das multas atuais "
         "× propostas? Qual a expectativa de arrecadação?")
rotulado("Emenda sugerida:", "teto de acumulação para multas diárias (ex.: 30 dias) e "
         "obrigatoriedade de divulgar anualmente a conversão UFESP→R$ das principais infrações.")

doc.add_heading("3.4 Medidas de força e devido processo", level=2)
rotulado("Dispositivos:", "recurso sem efeito suspensivo contra interdição e lacração (arts. "
         "319, §3º e 324, parágrafo único); demolição com execução subsidiária sem ordem "
         "judicial (art. 330, §2º); taxa de administração de 100% sobre custos de execução "
         "(arts. 101, §3º, 177, §2º, 326, 330 — hoje o padrão da Lei 3.531/68 é 20%); cassação "
         "com impedimento de 5 anos alcançando os sócios (art. 327, §2º — hoje 3 anos, restrito "
         "ao proprietário).")
rotulado("Impacto:", "concentração de poder coercitivo no Executivo com garantias reduzidas ao "
         "administrado; a taxa de 100% tem caráter punitivo cumulado à multa.")
rotulado("Pergunta ao Executivo:", "por que a supressão do efeito suspensivo mesmo fora de "
         "situações de risco iminente? Qual a justificativa técnica para elevar a taxa de "
         "administração de 20% para 100%?")
rotulado("Emenda sugerida:", "efeito suspensivo do recurso salvo risco iminente motivado; taxa "
         "de administração proporcional ao custo efetivamente comprovado; gradação do "
         "impedimento do art. 327, §2º conforme a gravidade.")

doc.add_heading("3.5 Segurança jurídica: revogações e ausência de transição", level=2)
rotulado("Dispositivos:", "art. 334 (lista de 23 normas) e ausência de disposições transitórias "
         "(arts. 331–335).")
rotulado("Impacto:", "dezenas de LCs que alteraram a Lei 3.531/68 — ex.: LC 450/2002, 533/2005, "
         "712/2011, 866/2014, 955/2017, 996/2018, 1105/2020, 1140/2021, 1189/2023, 1217/2023, "
         "1279/2024 — não constam da lista; com a revogação da lei-base, seu status fica "
         "indefinido [Verificar]. O PLC tampouco disciplina a transição de licenças vigentes, "
         "autos de infração lavrados e processos em curso.")
rotulado("Pergunta ao Executivo:", "qual o critério da lista do art. 334? As LCs alteradoras "
         "remanescem vigentes de forma autônoma? Como ficam autos e licenças em curso?")
rotulado("Emenda sugerida:", "revisar a cláusula revogatória (incluindo expressamente as LCs "
         "esgotadas) e acrescentar capítulo de disposições transitórias (validade das licenças "
         "emitidas; regime dos processos sancionatórios em andamento; vacatio legis adequada).")

doc.add_heading("3.6 Lacunas materiais a esclarecer", level=2)
bullet("Inflamáveis, explosivos e postos de combustível: capítulo suprimido sem regra "
       "substituta expressa; some a distância mínima de 200 m de escolas/hospitais para novos "
       "postos. [Verificar cobertura pela LUOS e por normas estaduais/federais]")
bullet("Cães na praia: caem as análises mensais obrigatórias da qualidade da areia com "
       "publicação no Diário Oficial (LC 1140/2021). [Verificar]")
bullet("Banho e tosa/hotéis de animais: guarda de imagens reduzida de 60 (LC 1217/2023) para 30 "
       "dias e hotéis/creches fora do rol — retrocesso pontual sobre norma de 2023. [Verificar]")
bullet("Placas de advertência sobre bebidas alcoólicas e energéticas (arts. 123-A/123-B atuais) "
       "desaparecem.")

doc.add_heading("3.7 Novas obrigações ao setor privado — custo regulatório", level=2)
rotulado("Dispositivos:", "Título III (drenagem): estudo técnico de impacto quando exigido "
         "(art. 120, I), corresponsabilidade de construtoras por 5 anos (art. 121), caixa de "
         "gordura com limpeza certificada e afixação de certificado (art. 115); Cap. IX do Tít. "
         "V (concessionárias).")
rotulado("Impacto:", "custos novos para construção civil, condomínios e comércio; risco de "
         "exigências abertas (“quando exigido”) sem critério legal.")
rotulado("Pergunta ao Executivo:", "qual a estimativa de custo de conformidade e o critério "
         "objetivo para exigir o estudo de impacto de drenagem?")
rotulado("Emenda sugerida:", "critérios objetivos (porte/área do empreendimento) para o estudo "
         "de impacto; dispensa das obrigações acessórias para pequenos geradores e MEI quando "
         "desproporcionais.")

doc.add_heading("3.8 Moldura principiológica", level=2)
rotulado("Dispositivos:", "arts. 2º e 3º (função socioambiental, direito à cidade, gestão "
         "democrática, precaução).")
rotulado("Impacto:", "linguagem do Estatuto da Cidade que pode fundamentar exigências abertas e "
         "interpretações extensivas contra o particular; como contrapeso, o próprio texto "
         "positiva a desburocratização (art. 2º, XI) e a livre iniciativa (art. 3º, XI), e o "
         "art. 332 veda analogias e interpretações extensivas.")
rotulado("Pergunta ao Executivo:", "confirmação de que os princípios dos arts. 2º–3º não criam, "
         "por si, obrigações autônomas exigíveis do particular.")
rotulado("Emenda sugerida:", "se necessário, regra interpretativa reforçando que obrigações "
         "decorrem apenas de previsão expressa em lei ou regulamento dela derivado.")

# ===================================================== 4. ENCAMINHAMENTOS
doc.add_heading("4. Encaminhamentos sugeridos à Comissão", level=1)
bullet("Oficiar o Executivo consolidando as perguntas das seções 3.2 a 3.8, com destaque para: "
       "minutas de decretos regulamentadores; estudo de impacto das multas; critério da lista "
       "do art. 334; regime de transição.")
bullet("Audiências públicas/técnicas: setor da construção civil e condomínios (Título III), "
       "concessionárias (Cap. IX do Tít. V), comércio e bares (sossego, licenciamento, "
       "sanções), entidades de proteção animal (arts. 204–214).")
bullet("Preparar bloco de emendas prioritárias: (i) disposições transitórias e cláusula "
       "revogatória; (ii) pisos legais mínimos nas matérias delegadas a decreto; (iii) teto de "
       "multas diárias e taxa de administração proporcional; (iv) efeito suspensivo de recurso "
       "fora de risco iminente; (v) critérios objetivos para o estudo de impacto de drenagem.")
bullet("Registrar no parecer o apoio da Comissão aos eixos de desburocratização (seção 3.1).")

# ================================================================ ANEXO
doc.add_page_break()
doc.add_heading("Anexo — Síntese quantitativa e referência", level=1)
p("Destino dos 640 artigos do código vigente (Lei nº 3.531/1968, consolidada):")
tabela(
    ("Destino no PLC", "Artigos"),
    (("Alterado (mérito)", 202), ("Condensado", 145), ("Remetido a outra norma", 142),
     ("Revogado sem correspondente", 82), ("Já revogado anteriormente", 49),
     ("Mantido/equivalente", 20), ("Total", 640)),
    (11.0, 4.0),
)
doc.add_paragraph()
p("Detalhamento completo — mapa por Título, mudanças substantivas, correspondência artigo a "
  "artigo, dispositivos novos e revogações expressas — na planilha comparativo_posturas.xlsx "
  "(mesma pasta deste relatório). Fontes: Lei nº 3.531/1968 consolidada (LeisMunicipais, "
  "06/01/2025) e PLC PA 59314/2025-78.")

doc.save(SAIDA)
print(f"OK — {SAIDA.name}")
