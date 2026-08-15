# -*- coding: utf-8 -*-
"""
Gera nota_inovacoes_emendas.docx — RESERVADO, uso interno do gabinete.
Catálogo de inovações regulatórias/de código de posturas NÃO incorporadas pelo
PLC PA 59314/2025-78, com emendas propostas, base legal e via de veiculação.
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
GOLD = RGBColor(0xC9, 0xA2, 0x4B)
VERM = RGBColor(0x9C, 0x27, 0x2A)
BASE = Path(__file__).parent
SAIDA = BASE / "nota_inovacoes_emendas.docx"

doc = Document()
_zoom = doc.settings.element.find(qn("w:zoom"))
if _zoom is not None and _zoom.get(qn("w:percent")) is None:
    _zoom.set(qn("w:percent"), "100")

sec = doc.sections[0]
sec.page_width, sec.page_height = Mm(210), Mm(297)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Cm(2.2))

normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.13
for nome, tam in (("Heading 1", 15), ("Heading 2", 12.5)):
    st = doc.styles[nome]
    st.font.name = "Arial"
    st.font.size = Pt(tam)
    st.font.bold = True
    st.font.color.rgb = NAVY

head_p = sec.header.paragraphs[0]
head_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = head_p.add_run("USO INTERNO — RESERVADO · Gabinete")
hr.font.size = Pt(8)
hr.font.color.rgb = VERM
hr.bold = True

foot_p = sec.footer.paragraphs[0]
foot_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = foot_p.add_run()
for tag, attrs, text in (("w:fldChar", {"w:fldCharType": "begin"}, None),
                         ("w:instrText", {"xml:space": "preserve"}, " PAGE "),
                         ("w:fldChar", {"w:fldCharType": "end"}, None)):
    el = OxmlElement(tag)
    for k, v in attrs.items():
        el.set(qn(k), v)
    if text:
        el.text = text
    run._r.append(el)
run.font.size = Pt(9)


def p(texto, bold=False, italic=False, size=None, align=None, color=None):
    par = doc.add_paragraph()
    r = par.add_run(texto)
    r.bold, r.italic = bold, italic
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    if align:
        par.alignment = align
    return par


def bullet(texto):
    par = doc.add_paragraph(style="List Bullet")
    par.add_run(texto)
    return par


def rotulado(rotulo, texto):
    par = doc.add_paragraph()
    r = par.add_run(rotulo + " ")
    r.bold = True
    r.font.color.rgb = NAVY
    par.add_run(texto)
    return par


def caixa(texto):
    par = doc.add_paragraph()
    r = par.add_run(texto)
    r.italic = True
    par.paragraph_format.left_indent = Cm(0.5)
    pPr = par._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "C9A24B")
    pbdr.append(left)
    ref = None
    for tag in ("w:spacing", "w:ind", "w:contextualSpacing", "w:jc", "w:rPr"):
        el = pPr.find(qn(tag))
        if el is not None:
            ref = el
            break
    (ref.addprevious(pbdr) if ref is not None else pPr.append(pbdr))
    return par


def tabela(cabecalho, linhas, larguras_cm, fs=9):
    t = doc.add_table(rows=1, cols=len(cabecalho))
    t.style = "Table Grid"
    for i, (h, w) in enumerate(zip(cabecalho, larguras_cm)):
        cel = t.rows[0].cells[i]
        cel.width = Cm(w)
        r = cel.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(fs)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "1B2A4A")
        cel._tc.get_or_add_tcPr().append(shd)
    for linha in linhas:
        cells = t.add_row().cells
        for i, (v, w) in enumerate(zip(linha, larguras_cm)):
            cells[i].width = Cm(w)
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(fs)
    return t


# ================================================================ cabeçalho
p("NOTA INTERNA — RESERVADA", bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER, color=VERM)
p("Uso exclusivo do gabinete", italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=VERM)
doc.add_paragraph()
p("Inovações que o PLC deixou de fora — catálogo de emendas modernizadoras",
  bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, color=NAVY)
p("PLC PA 59314/2025-78 (novo Código de Posturas) — o que se pode inserir via comissão ou "
  "emenda individual", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, color=GOLD)
p("Câmara Municipal de Santos — julho de 2026", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

# ================================================================= objetivo
doc.add_heading("1. Objetivo", level=1)
p("O PLC da Prefeitura limpa o código de 1968, mas não incorpora o instrumental de "
  "modernização pró-liberdade consolidado no direito brasileiro desde 2019. Esta nota reúne as "
  "inovações ausentes, cada uma com a lacuna correspondente no PLC, a emenda proposta, a base "
  "legal/exemplo e a via de veiculação. Prioridade nos blocos 2 e 3 (liberdade econômica e "
  "governo digital), que são a nossa agenda; os blocos 4 e 5 são menu, com os trade-offs "
  "sinalizados.")
caixa("Nota de estratégia: você já tem PL de liberdade econômica/baixo risco em tramitação "
      "[inserir nº do PL]. A emenda do bloco 2 deve REMETER e HARMONIZAR o Código a esse PL — "
      "não duplicar —, transformando o Código de Posturas no veículo que operacionaliza a sua "
      "proposta. Isso amarra as duas iniciativas sob a sua autoria.")

# ============================================ bloco 2 — liberdade econômica
doc.add_heading("2. Bloco prioritário — Liberdade econômica e licenciamento", level=1)

doc.add_heading("2.1 Dispensa de licenciamento e alvará para atividades de baixo risco", level=2)
rotulado("Inovação:", "atividade de baixo risco opera com o CNPJ, sem alvará nem licença "
         "prévios; nível intermediário obtém liberação automática após o registro. "
         "Autodeclaração 100% digital (modelo Porto Alegre).")
rotulado("Base legal:", "Lei federal 13.874/2019 (Liberdade Econômica), art. 3º, I e VI; "
         "Resoluções CGSIM nº 51/2019 e 57/2020 (definem baixo risco A/B; a norma municipal "
         "própria afasta a federal). Exemplos: São Paulo (~911 atividades dispensadas), Minas "
         "Gerais (915), Belo Horizonte (275 por decreto), Balneário Camboriú (1.114).")
rotulado("Lacuna no PLC:", "mantém licença prévia para todos (art. 221); apenas cita “alto "
         "risco pela legislação municipal” (art. 222, §1º), sem instituir a dispensa de baixo "
         "risco.")
rotulado("Emenda:", "capítulo remetendo à classificação de risco e à dispensa nos termos do PL "
         "de liberdade econômica em tramitação, com autodeclaração eletrônica e lista mínima de "
         "baixo risco; vedação de exigir alvará onde a lei dispensa.")

doc.add_heading("2.2 Silêncio positivo (aprovação tácita)", level=2)
rotulado("Inovação:", "vencido o prazo de análise sem decisão, o pedido é deferido "
         "automaticamente — inverte a lógica do silêncio negativo.")
rotulado("Base legal:", "Lei 13.874/2019, art. 3º (direito à análise em prazo, sob pena de "
         "aprovação tácita).")
rotulado("Lacuna no PLC:", "o art. 222, §1º só permite “operar exibindo o protocolo” — herança "
         "literal de 1968 (art. 428, §1º) — sem deferimento tácito nem consequência à inércia "
         "do Executivo.")
rotulado("Emenda:", "findo o prazo de 15 dias sem decisão fundamentada, a licença considera-se "
         "deferida para todos os fins, ressalvado o alto risco.")

doc.add_heading("2.3 Análise de Impacto Regulatório (AIR) e revisão do estoque", level=2)
rotulado("Inovação:", "toda nova obrigação passa por AIR (custo/benefício, alternativas); o "
         "estoque de exigências é revisto periodicamente, com cláusula de caducidade (sunset) "
         "para regras que não se justifiquem.")
rotulado("Base legal:", "Decreto federal 10.411/2020 (AIR — referência a adaptar ao município, "
         "pois vincula a União); boas práticas regulatórias (OCDE).")
rotulado("Lacuna no PLC:", "delega amplamente a decreto (praias, calçadas, ambulantes, "
         "dosimetria) sem exigir AIR nem revisão periódica.")
rotulado("Emenda:", "condicionar a criação de obrigações por decreto à AIR prévia e instituir "
         "revisão quinquenal do estoque de exigências, com sunset para as não confirmadas.")

# ============================================ bloco 3 — governo digital
doc.add_heading("3. Bloco prioritário — Governo digital e desburocratização", level=1)

doc.add_heading("3.1 “Solicita uma única vez”, interoperabilidade e alvará sem papel", level=2)
rotulado("Inovação:", "o poder público não pode exigir documento/informação que já possui ou "
         "que outro órgão detém (once-only); processo e alvará totalmente eletrônicos, sem via "
         "física.")
rotulado("Base legal:", "Lei federal 14.129/2021 (Governo Digital), princípios de solicitação "
         "única, interoperabilidade e eliminação de exigências desnecessárias.")
rotulado("Lacuna no PLC:", "o art. 224, §1º ainda obriga conservar o alvará afixado impresso — "
         "o oposto de digital; o texto não veda a reexigência de documentos.")
rotulado("Emenda:", "alvará e processos 100% digitais com dispensa da via física; vedação "
         "expressa de exigir documento já em poder do Município; QR-code de verificação em vez "
         "de afixação.")

# ============================================ bloco 4 — temas contemporâneos
doc.add_heading("4. Bloco temático — questões urbanas atuais (menu)", level=1)

doc.add_heading("4.1 Regime de convivência da economia noturna", level=2)
rotulado("Inovação:", "polos gastronômicos/de entretenimento com horário e limites próprios; "
         "mediação de vizinhança e “acordo de boa vizinhança”; escalonamento educativo antes da "
         "multa — trata a noite como política de convivência e economia, não só como caso de "
         "polícia.")
rotulado("Base legal:", "experiências nacionais (debates em BH, Curitiba) e internacionais "
         "(Londres, Amsterdã, Berlim — “gestão da noite”); princípio da proporcionalidade da "
         "LLE.")
rotulado("Lacuna no PLC:", "endurece a sanção sonora (apreensão de equipamentos, arts. 137–143) "
         "sem instrumentos de mediação nem regime de polo.")
rotulado("Emenda:", "prever zonas/polos com regime próprio, mediação prévia e advertência antes "
         "da apreensão para o primeiro evento — pró-negócio e pró-sossego.")

doc.add_heading("4.2 Conformidade LGPD do videomonitoramento obrigatório", level=2)
rotulado("Lacuna/Emenda:", "o art. 212 obriga câmeras em banho e tosa (guarda de imagens) sem "
         "uma linha de proteção de dados. Emenda de conformidade à LGPD (Lei 13.709/2018): "
         "finalidade, controle de acesso, prazo e direitos do titular. Fecha o flanco de "
         "modernização e o de liberdades.")

doc.add_heading("4.3 Micromobilidade (patinetes e bicicletas compartilhadas)", level=2)
rotulado("Lacuna/Emenda:", "o PLC só toca de leve (art. 175, “pontos de aluguel”). Emenda "
         "disciplinando estacionamento ordenado, áreas de exclusão (geofencing), remoção de "
         "equipamentos abandonados e responsabilidade do operador — problema urbano atual não "
         "endereçado.")

doc.add_heading("4.4 Poluição luminosa [expansão regulatória — avaliar]", level=2)
rotulado("Lacuna/Emenda:", "ausente. Possível inclusão de diretrizes (luz direcionada para "
         "baixo, controle de ofuscamento) inspiradas no PL federal 1.400/2021 e na ABNT NBR "
         "5101/2024. Alerta do mandato: é criação de nova obrigação — pesar custo/benefício "
         "antes de propor; talvez melhor como diretriz, não como exigência sancionável.")

doc.add_heading("4.5 Calçadas — requalificação em vias estruturais [trade-off fiscal]", level=2)
rotulado("Lacuna/Emenda:", "o PLC mantém toda a responsabilidade no proprietário (art. 166). "
         "Modelo alternativo: programa municipal de requalificação em vias estruturais (ex.: "
         "“Passeio Livre”/PEC, São Paulo), desonerando o cidadão nos eixos estratégicos. Alerta "
         "do mandato: impacto fiscal é o critério nº 1 — só propor com fonte de custeio clara "
         "e recorte restrito; caso contrário, manter no proprietário com padrão simplificado.")

doc.add_heading("4.6 Parklets e urbanismo tático — diretrizes pró-uso", level=2)
rotulado("Lacuna/Emenda:", "o PLC prevê parklet mas remete a regulamento (art. 176). Emenda "
         "fixando diretrizes pró-uso (prazo célere, isenção/desconto para pequeno negócio, "
         "padrão simplificado) para não morrer no decreto.")

# ============================================ bloco 5 — conformidade cooperativa
doc.add_heading("5. Conformidade cooperativa (reforço)", level=1)
p("O PLC já traz advertência (art. 304) e termo de compromisso (arts. 293–298). Emenda de "
  "reforço: tornar a advertência/orientação a regra para a primeira infração leve de MEI/ME/EPP "
  "e microempreendedor, com multa apenas na reincidência — coerente com o art. 55 da LLE "
  "(dupla visita orientadora) e com a defesa do pequeno negócio.")

# ============================================ quadro-síntese
doc.add_heading("6. Quadro-síntese das emendas", level=1)
tabela(
    ("#", "Emenda", "Lacuna no PLC", "Base legal / exemplo"),
    (
        ("1", "Dispensa de licença/alvará p/ baixo risco + autodeclaração digital",
         "Licença prévia p/ todos (art. 221)",
         "Lei 13.874/2019; CGSIM 51/2019 e 57/2020; SP, MG, BH"),
        ("2", "Silêncio positivo (deferimento tácito ao fim do prazo)",
         "Só “operar com protocolo” (art. 222 §1º)", "Lei 13.874/2019, art. 3º"),
        ("3", "AIR p/ nova obrigação + revisão quinquenal (sunset)",
         "Delega a decreto sem AIR", "Decreto 10.411/2020 (adaptar)"),
        ("4", "Once-only + interoperabilidade + alvará sem papel",
         "Alvará afixado impresso (art. 224 §1º)", "Lei 14.129/2021"),
        ("5", "Conformidade LGPD do videomonitoramento obrigatório",
         "Câmeras sem proteção de dados (art. 212)", "Lei 13.709/2018 (LGPD)"),
        ("6", "Regime de polo noturno + mediação + advertência antes da apreensão",
         "Só endurece sanção sonora (137–143)", "Proporcionalidade (LLE); experiências"),
        ("7", "Micromobilidade: geofencing, estacionamento, responsabilidade do operador",
         "Trata de leve (art. 175)", "—"),
        ("8", "Poluição luminosa (diretriz) [avaliar custo]",
         "Ausente", "PL fed. 1.400/2021; ABNT NBR 5101/2024"),
        ("9", "Requalificação de calçadas em vias estruturais [trade-off fiscal]",
         "Tudo no proprietário (art. 166)", "SP Passeio Livre / PEC"),
        ("10", "Diretrizes pró-uso p/ parklets",
         "Remetido a decreto (art. 176)", "SP parklets"),
        ("11", "Dupla visita orientadora antes de multar MEI/ME/EPP",
         "Multa direta possível", "Lei 13.874/2019, art. 55"),
    ),
    (0.7, 5.3, 4.3, 5.7),
)

# ============================================ via de veiculação
doc.add_heading("7. Via de veiculação", level=1)
bullet("Substitutivo / emenda coletiva da comissão especial — para os blocos estruturantes que "
       "reorganizam o texto (emendas 1 a 4). É a forma mais robusta e com autoria da comissão.")
bullet("Emenda individual aditiva/modificativa — para as inserções pontuais (emendas 5 a 11), "
       "que não dependem de reescrever capítulos e podem levar a sua assinatura.")
bullet("Sequenciamento sugerido: priorizar 1–5 (liberdade econômica, governo digital, LGPD), "
       "que são consensuais e de alto retorno político; tratar 6–11 como menu conforme o "
       "apetite da comissão e o custo fiscal.")

# ============================================ fontes
doc.add_heading("8. Base normativa e fontes", level=1)
for t in (
    "Lei federal 13.874/2019 — Liberdade Econômica (dispensa de baixo risco; aprovação tácita; "
    "dupla visita, art. 55).",
    "Resoluções CGSIM nº 51/2019 e 57/2020 — definição de baixo risco (níveis A e B).",
    "Lei federal 14.129/2021 — Governo Digital (solicitação única; interoperabilidade).",
    "Decreto federal 10.411/2020 — Análise de Impacto Regulatório (referência a adaptar).",
    "Lei federal 13.709/2018 — LGPD.",
    "ABNT NBR 5101/2024 — iluminação pública; PL federal 1.400/2021 — poluição luminosa.",
    "Exemplos municipais: São Paulo (dispensa ~911 atividades; Programa Passeio Livre/PEC; "
    "parklets), Belo Horizonte (Decreto — 275 atividades), Minas Gerais (915), Balneário "
    "Camboriú (classificação 1.114/126/173).",
):
    bullet(t)
p("")
p("Documento de trabalho interno. Itens marcados [avaliar]/[trade-off] envolvem criação de "
  "obrigação ou despesa e devem passar pelo crivo fiscal do mandato antes de propostos. "
  "Número do PL de liberdade econômica de sua autoria a inserir no item 2.1.",
  italic=True, size=9)

doc.save(SAIDA)
print(f"OK — {SAIDA.name}")
