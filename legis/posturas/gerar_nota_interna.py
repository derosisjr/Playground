# -*- coding: utf-8 -*-
"""
Gera nota_interna_gabinete.docx — documento RESERVADO, de uso interno do gabinete.
Tese: o PLC do Código de Posturas faz faxina (remove anacronismos), mas não
moderniza sob liberdade econômica/desburocratização. Análise sob o perfil do
mandato (oposição/PL) + quadro de emendas modernizadoras.
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
GOLD = RGBColor(0xC9, 0xA2, 0x4B)
VERM = RGBColor(0x9C, 0x27, 0x2A)
BASE = Path(__file__).parent
SAIDA = BASE / "nota_interna_gabinete.docx"

doc = Document()
_zoom = doc.settings.element.find(qn("w:zoom"))
if _zoom is not None and _zoom.get(qn("w:percent")) is None:
    _zoom.set(qn("w:percent"), "100")

sec = doc.sections[0]
sec.page_width, sec.page_height = Mm(210), Mm(297)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Cm(2.5))

normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15
for nome, tam in (("Heading 1", 15), ("Heading 2", 12.5)):
    st = doc.styles[nome]
    st.font.name = "Arial"
    st.font.size = Pt(tam)
    st.font.bold = True
    st.font.color.rgb = NAVY

# cabeçalho e rodapé com marcação de reserva
head_p = sec.header.paragraphs[0]
head_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = head_p.add_run("USO INTERNO — RESERVADO · Gabinete")
hr.font.size = Pt(8)
hr.font.color.rgb = VERM
hr.bold = True

foot_p = sec.footer.paragraphs[0]
foot_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = foot_p.add_run()
for tag, attrs, text in (("w:fldChar", {"w:fldCharType": "begin"}, None),
                         ("w:instrText", {"xml:space": "preserve"}, " PAGE "),
                         ("w:fldChar", {"w:fldCharType": "end"}, None)):
    el = OxmlElement(tag)
    for k, v in attrs.items():
        el.set(qn(k), v)
    if text:
        el.text = text
    run._r.append(el)
run.font.size = Pt(9)


def p(texto, bold=False, italic=False, size=None, align=None, color=None):
    par = doc.add_paragraph()
    r = par.add_run(texto)
    r.bold, r.italic = bold, italic
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    if align:
        par.alignment = align
    return par


def bullet(texto):
    par = doc.add_paragraph(style="List Bullet")
    par.add_run(texto)
    return par


def rotulado(rotulo, texto):
    par = doc.add_paragraph()
    r = par.add_run(rotulo + " ")
    r.bold = True
    r.font.color.rgb = NAVY
    par.add_run(texto)
    return par


def caixa(texto):
    par = doc.add_paragraph()
    r = par.add_run(texto)
    r.italic = True
    r.font.size = Pt(11)
    par.paragraph_format.left_indent = Cm(0.5)
    pPr = par._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "C9A24B")
    pbdr.append(left)
    # pBdr deve preceder spacing/ind/jc/rPr no pPr (ordem do schema)
    ref = None
    for tag in ("w:spacing", "w:ind", "w:contextualSpacing", "w:jc", "w:rPr"):
        el = pPr.find(qn(tag))
        if el is not None:
            ref = el
            break
    if ref is not None:
        ref.addprevious(pbdr)
    else:
        pPr.append(pbdr)
    return par


def tabela(cabecalho, linhas, larguras_cm):
    t = doc.add_table(rows=1, cols=len(cabecalho))
    t.style = "Table Grid"
    for i, (h, w) in enumerate(zip(cabecalho, larguras_cm)):
        cel = t.rows[0].cells[i]
        cel.width = Cm(w)
        r = cel.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "1B2A4A")
        cel._tc.get_or_add_tcPr().append(shd)
    for linha in linhas:
        cells = t.add_row().cells
        for i, (v, w) in enumerate(zip(linha, larguras_cm)):
            cells[i].width = Cm(w)
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(9.5)
    return t


# ============================================================== cabeçalho
p("NOTA INTERNA — RESERVADA", bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER, color=VERM)
p("Uso exclusivo do gabinete — não circular na comissão nem com o Executivo",
  italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=VERM)
doc.add_paragraph()
p("Faxina × Modernização: o que o PLC do Código de Posturas deixou de fazer",
  bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, color=NAVY)
p("PLC PA 59314/2025-78 — leitura sob o mandato (defesa do contribuinte, liberdade "
  "econômica, desburocratização, fiscalização do Executivo)",
  size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, color=GOLD)
p("Câmara Municipal de Santos — julho de 2026", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

# ================================================================ veredito
doc.add_heading("1. Tese e veredito", level=1)
caixa("O PLC faz faxina, não reforma. Ele remove o entulho de 1968 sem instalar a fiação do "
      "século XXI. É contemporâneo onde amplia o poder de polícia e as obrigações privadas; é "
      "meramente janitorial onde poderia libertar a economia. Chamamos isso de modernização "
      "assimétrica.")
p("A impressão da presidência da comissão está correta no essencial. Dois ajustes de precisão:")
bullet("Liberdade econômica e desburocratização: é faxina, não modernização — plenamente "
       "confirmado (seção 2 e 3).")
bullet("Problemas atuais da cidade: há, sim, conteúdo novo (drenagem, cabeamento, zoonoses, "
       "bem-estar animal), mas quase todo na direção de mais Estado — o que reforça a nossa "
       "crítica em vez de contrariá-la (seção 4).")

# =========================================================== liberdade
doc.add_heading("2. Liberdade econômica: faxina, não modernização", level=1)
p("Modernizar licenciamento tem hoje significado técnico preciso, dado pela Lei federal de "
  "Liberdade Econômica (13.874/2019) e pela agenda Redesim/CGSIM: presunção de boa-fé do "
  "particular; dispensa de alvará para atividades de baixo risco; silêncio positivo (aprovação "
  "tácita); análise de impacto regulatório; autodeclaração com fiscalização a posteriori. "
  "O PLC não importa nada disso.")
bullet("Mantém intocado o paradigma de 1968 — licença prévia para todos (art. 221). Muda a "
       "decoração, não o mecanismo.")
bullet("Vende como agilidade o “pode começar com o protocolo” (art. 222, §1º), mas isso já "
       "existe desde 1968 (art. 428, §1º do código atual) — e o PLC ainda acrescenta exigências "
       "para iniciar (AVCB + Licença Ambiental para alto risco). Aqui é mais exigente, não menos.")
bullet("Declara “livre iniciativa” e “desburocratização” como princípios (arts. 2º, XI e 3º, "
       "XI), mas não os operacionaliza em nenhum direito autoexecutável. É modernização de "
       "vocabulário, não de mecanismo.")
rotulado("[Verificar]:", "o art. 222, §1º cita “atividades de alto risco pela legislação "
         "municipal”, o que sugere já existir norma de classificação de risco em Santos. Se "
         "existe, o PLC perdeu a chance de consolidá-la e reafirmar a presunção de liberdade; "
         "se não existe, perdeu a chance de criá-la. Nos dois casos, o Código em si não "
         "moderniza o licenciamento.")

# ====================================================== desburocratização
doc.add_heading("3. A desburocratização ilusória", level=1)
p("Cortar 48% dos artigos não é cortar 48% da burocracia. Burocracia se mede em passos, "
  "licenças, documentos e prazos que uma pessoa real enfrenta — não em contagem de artigos. "
  "E parte relevante do enxugamento é realocação, não redução:")
bullet("Matérias saíram do texto legal para “decreto/regulamento” e “legislação específica” "
       "(praias, calçadas, ambulantes, dosimetria). A exigência não sumiu — migrou para fora do "
       "controle da Câmara e da vista do cidadão. Isso é menos transparência, não menos "
       "burocracia.")
bullet("Onde inovou, muitas vezes regulou mais: estudo de impacto de drenagem, "
       "corresponsabilidade de 5 anos das construtoras, caixa de gordura certificada e afixada, "
       "câmeras obrigatórias, CCT com ARTs, declarações para eventos.")
caixa("Resumo: onde encurtou, muitas vezes delegou; onde inovou, muitas vezes regulou mais.")

# ====================================================== assimetria
doc.add_heading("4. Modernização assimétrica: quando inovou, inovou para mais Estado", level=1)
p("Sendo rigorosos (a análise pede rigor, não militância): há conteúdo novo que encara "
  "problemas atuais — não é só remoção de anacronismo:")
bullet("Título de Drenagem Urbana (arts. 105–127) — Santos convive com alagamento crônico.")
bullet("Cabeamento aéreo e concessionárias (arts. 263–279) — poluição visual de fios.")
bullet("Zoonoses, bem-estar animal, cães na praia, fim de charretes — pautas atuais.")
bullet("Fiscalização eletrônica e termo de compromisso — modernização processual real.")
p("Mas repare a direção: quase todo esse “novo” chega em forma de nova obrigação e nova multa "
  "(drenagem até 20.000 UFESP; câmeras; corresponsabilidade). O PLC moderniza na direção de "
  "mais Estado, não de mais liberdade — exatamente a modernização que não nos interessa.")
p("Para ser justo, há duas liberalizações genuínas, a exceção que confirma a regra: o fim da "
  "tabela de horário por ramo (art. 229) — liberdade econômica de verdade — e o tratamento "
  "diferenciado a MEI/ME/EPP na dosimetria (art. 300, V).")

# ====================================================== emendas
doc.add_heading("5. Quadro de emendas modernizadoras (o que faltou)", level=1)
p("Transforma a crítica em proposta — espaço para a oposição deixar digital própria, "
  "preenchendo o vácuo pró-liberdade que o Executivo não ocupou.")
tabela(
    ("Nº", "O que falta no PLC", "Emenda modernizadora", "Base / referência"),
    (
        ("1", "Paradigma de licença prévia para todos, sem presunção de liberdade",
         "Ancorar o Código na Lei de Liberdade Econômica: boa-fé do particular, dispensa de "
         "licença/alvará para baixo risco, silêncio positivo, AIR obrigatória para nova "
         "obrigação criada por decreto",
         "Lei 13.874/2019; Redesim/CGSIM"),
        ("2", "Autorização prévia como regra",
         "Licenciamento por autodeclaração com fiscalização a posteriori — verificar depois, "
         "não autorizar antes",
         "LLE, art. 3º"),
        ("3", "Via física do alvará ainda exigida (art. 224, §1º)",
         "Digitalização real: processo eletrônico, assinatura digital e dispensa da via física",
         "Gov digital"),
        ("4", "Reexigência de documentos que o Município já tem",
         "Princípio once-only: vedar exigir documento já em poder do poder público",
         "Lei 14.129/2021 (Gov. Digital)"),
        ("5", "Exigências sem prazo de validade / revisão",
         "Cláusula de revisão periódica (sunset) das exigências e das obrigações delegadas a "
         "decreto",
         "Boas práticas regulatórias"),
        ("6", "Videovigilância privada obrigatória sem proteção de dados (art. 212)",
         "Cláusula de conformidade à LGPD: finalidade, acesso, retenção e direitos do titular "
         "das imagens",
         "Lei 13.709/2018 (LGPD)"),
        ("7", "Multa diária sem teto e taxa de administração de 100%",
         "Teto de acumulação da multa diária (ex.: 30 dias) e taxa proporcional ao custo "
         "comprovado",
         "Proporcionalidade"),
    ),
    (1.0, 5.0, 6.6, 3.4),
)

# ====================================================== LGPD destaque
doc.add_heading("6. Destaque: a lacuna de LGPD", level=1)
p("O PLC obriga videovigilância privada (câmeras em banho e tosa, art. 212, com guarda de "
  "imagens) sem uma linha sobre proteção de dados — finalidade, controle de acesso, prazo e "
  "direitos do titular. É, ao mesmo tempo, atraso de modernização e risco às liberdades "
  "individuais: o Município manda o particular vigiar e guardar imagens, mas não o obriga a "
  "proteger esses dados. Emenda de conformidade à LGPD resolve os dois flancos e é bandeira "
  "boa de defender publicamente.")

# ====================================================== enquadramento
doc.add_heading("7. Enquadramento político e fala ao eleitor", level=1)
p("A leitura muda a missão da comissão: não é aprovar ou rejeitar uma faxina — a faxina é "
  "bem-vinda. É usar o vácuo pró-liberdade que o Executivo deixou como espaço para a oposição "
  "inserir a reforma que interessa ao contribuinte. Nunca enquadrar como “aprovei porque o "
  "prefeito propôs”, e sim “a oposição trouxe o século XXI que faltava”.")
p("Rascunhos de fala:", bold=True)
caixa("“O governo limpou a poeira do código de 1968; nós trouxemos o século XXI: menos licença "
      "para o pequeno negócio, aprovação automática quando o prazo vence e fim da exigência de "
      "papel que a própria prefeitura já tem.”")
caixa("“Desregulamentar de verdade é dispensar alvará de baixo risco e confiar no cidadão de "
      "boa-fé — não trocar seis regras velhas por uma multa nova.”")

# ====================================================== diligências
doc.add_heading("8. Diligências pendentes [Verificar]", level=1)
tabela(
    ("Item a verificar", "Por que importa"),
    (
        ("Santos já possui norma municipal de classificação de risco / dispensa de baixo risco "
         "(Redesim/CGSIM)?",
         "Calibra a emenda 1: consolidar o que existe ou criar do zero"),
        ("Valor da UFESP vigente na tramitação",
         "Converter as multas em reais e dimensionar o impacto ao contribuinte"),
        ("Cobertura de inflamáveis/postos de combustível por LUOS ou norma estadual",
         "Capítulo suprimido; some a distância mínima de 200 m de escolas/hospitais"),
        ("Status das LCs alteradoras não listadas no art. 334",
         "Segurança jurídica após a revogação da lei-base"),
    ),
    (8.0, 8.0),
)
p("")
p("Documento de trabalho interno. Baseado no texto do PLC e da Lei 3.531/1968 e nos "
  "entregáveis comparativo_posturas.xlsx e relatorio_posturas.docx. Itens não confirmados em "
  "fonte primária estão marcados [Verificar].", italic=True, size=9)

doc.save(SAIDA)
print(f"OK — {SAIDA.name}")
