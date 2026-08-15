# -*- coding: utf-8 -*-
"""
Gera dossie_gabinete.docx — DOSSIÊ ÚNICO, RESERVADO, uso interno do gabinete.
Consolida: panorama do PLC, leitura estratégica (perfil do mandato), análises
temáticas (delegação, multas, interferência estatal, outras), catálogo de
inovações não incorporadas e o TEXTO DAS EMENDAS pronto para protocolo.
PLC PA 59314/2025-78 (novo Código de Posturas de Santos) × Lei 3.531/1968.
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
GOLD = RGBColor(0xC9, 0xA2, 0x4B)
VERM = RGBColor(0x9C, 0x27, 0x2A)
BASE = Path(__file__).parent
SAIDA = BASE / "dossie_gabinete.docx"

doc = Document()
_zoom = doc.settings.element.find(qn("w:zoom"))
if _zoom is not None and _zoom.get(qn("w:percent")) is None:
    _zoom.set(qn("w:percent"), "100")

sec = doc.sections[0]
sec.page_width, sec.page_height = Mm(210), Mm(297)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Cm(2.2))

normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.12
for nome, tam in (("Heading 1", 15), ("Heading 2", 12), ("Heading 3", 10.5)):
    st = doc.styles[nome]
    st.font.name = "Arial"
    st.font.size = Pt(tam)
    st.font.bold = True
    st.font.color.rgb = NAVY if nome != "Heading 3" else GOLD

head_p = sec.header.paragraphs[0]
head_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = head_p.add_run("USO INTERNO — RESERVADO · Gabinete")
hr.font.size = Pt(8)
hr.font.color.rgb = VERM
hr.bold = True

foot_p = sec.footer.paragraphs[0]
foot_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = foot_p.add_run()
for tag, attrs, text in (("w:fldChar", {"w:fldCharType": "begin"}, None),
                         ("w:instrText", {"xml:space": "preserve"}, " PAGE "),
                         ("w:fldChar", {"w:fldCharType": "end"}, None)):
    el = OxmlElement(tag)
    for k, v in attrs.items():
        el.set(qn(k), v)
    if text:
        el.text = text
    run._r.append(el)
run.font.size = Pt(9)


def p(texto, bold=False, italic=False, size=None, align=None, color=None):
    par = doc.add_paragraph()
    r = par.add_run(texto)
    r.bold, r.italic = bold, italic
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    if align:
        par.alignment = align
    return par


def bullet(texto):
    par = doc.add_paragraph(style="List Bullet")
    par.add_run(texto)
    return par


def rotulado(rotulo, texto):
    par = doc.add_paragraph()
    r = par.add_run(rotulo + " ")
    r.bold = True
    r.font.color.rgb = NAVY
    par.add_run(texto)
    return par


def _borda(par, cor, largura="18"):
    pPr = par._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), largura)
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), cor)
    pbdr.append(left)
    ref = None
    for tag in ("w:spacing", "w:ind", "w:contextualSpacing", "w:jc", "w:rPr"):
        el = pPr.find(qn(tag))
        if el is not None:
            ref = el
            break
    (ref.addprevious(pbdr) if ref is not None else pPr.append(pbdr))


def caixa(texto, cor="C9A24B"):
    par = doc.add_paragraph()
    r = par.add_run(texto)
    r.italic = True
    par.paragraph_format.left_indent = Cm(0.5)
    _borda(par, cor)
    return par


def texto_legal(linhas):
    for ln in linhas:
        par = doc.add_paragraph()
        r = par.add_run(ln)
        r.font.size = Pt(10)
        par.paragraph_format.left_indent = Cm(0.6)
        par.paragraph_format.space_after = Pt(3)
        _borda(par, "1B2A4A", "12")


def tabela(cabecalho, linhas, larguras_cm, fs=9):
    t = doc.add_table(rows=1, cols=len(cabecalho))
    t.style = "Table Grid"
    for i, (h, w) in enumerate(zip(cabecalho, larguras_cm)):
        cel = t.rows[0].cells[i]
        cel.width = Cm(w)
        r = cel.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(fs)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "1B2A4A")
        cel._tc.get_or_add_tcPr().append(shd)
    for linha in linhas:
        cells = t.add_row().cells
        for i, (v, w) in enumerate(zip(linha, larguras_cm)):
            cells[i].width = Cm(w)
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(fs)
    return t


def emenda(num, tipo, disp, corpo, justif):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(8)
    r = par.add_run(f"EMENDA Nº {num} — {tipo}")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = VERM
    rotulado("Dispositivo:", disp)
    p("Texto proposto:", bold=True, size=10)
    texto_legal(corpo)
    rotulado("Justificativa:", justif)


# ==================================================================== CAPA
for _ in range(4):
    doc.add_paragraph()
p("DOSSIÊ DO GABINETE — RESERVADO", bold=True, size=22, align=WD_ALIGN_PARAGRAPH.CENTER, color=VERM)
p("Uso exclusivo do gabinete — não circular na comissão nem com o Executivo",
  italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=VERM)
doc.add_paragraph()
p("Novo Código de Posturas de Santos — análise integral e emendas",
  bold=True, size=17, align=WD_ALIGN_PARAGRAPH.CENTER, color=NAVY)
p("PLC PA 59314/2025-78  ×  Lei nº 3.531/1968 (texto consolidado)",
  size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=GOLD)
for _ in range(2):
    doc.add_paragraph()
p("Comissão Especial de Modernização do Código de Posturas", bold=True, size=11,
  align=WD_ALIGN_PARAGRAPH.CENTER)
p("Câmara Municipal de Santos — julho de 2026", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(3):
    doc.add_paragraph()
caixa("PENDÊNCIAS A PREENCHER: (i) número do PL de liberdade econômica de sua autoria, em "
      "tramitação, a citar na Emenda nº 1; (ii) valor da UFESP vigente, para converter as multas "
      "em reais. Itens não confirmados em fonte primária estão marcados [Verificar].", "9C272A")
doc.add_page_break()

# ================================================================= SUMÁRIO
doc.add_heading("Sumário", level=1)
for item in (
    "Parte I — Panorama: o que o PLC muda (faxina × modernização)",
    "Parte II — Leitura estratégica sob o mandato (apoiar / atacar / condicionar)",
    "Parte III — Análises temáticas: delegação a decreto; arquitetura das multas; "
    "interferência estatal; outras",
    "Parte IV — Inovações não incorporadas pelo PLC",
    "Parte V — Emendas (texto pronto para protocolo)",
    "Parte VI — Estratégia de veiculação e sequenciamento",
    "Parte VII — Diligências pendentes e fontes",
):
    bullet(item)
doc.add_page_break()

# ============================================================ PARTE I
doc.add_heading("Parte I — Panorama: o que o PLC muda", level=1)
caixa("Tese central: o PLC faz faxina, não reforma. Remove o entulho de 1968 sem instalar a "
      "fiação do século XXI. É contemporâneo onde amplia o poder de polícia e as obrigações "
      "privadas; é meramente janitorial onde poderia libertar a economia — modernização "
      "assimétrica.")
p("O PLC substitui ~640 artigos consolidados por 335, revogando nominalmente 23 normas (art. "
  "334). Destino dos 640 artigos vigentes: 202 alterados no mérito, 145 condensados, 142 "
  "remetidos a outras normas, 82 revogados sem correspondente, 49 já revogados antes do PLC e "
  "20 mantidos (detalhe na planilha comparativo_posturas.xlsx).")
doc.add_heading("O que efetivamente muda", level=2)
bullet("Nova moldura de princípios e direitos (arts. 1º–10), incluindo livre iniciativa e "
       "desburocratização, mas também linguagem do Estatuto da Cidade.")
bullet("Higiene condensada; controle de zoonoses, coleta seletiva e poluição ampliada (Tít. II).")
bullet("NOVO Título de Drenagem Urbana (arts. 105–127), com obrigações a imóveis, construtoras "
       "(corresponsabilidade de 5 anos) e concessionárias.")
bullet("Bem-estar público reorganizado; horário de funcionamento desregulamentado (art. 229); "
       "comércio ambulante e bancas remetidos a legislação específica.")
bullet("Remoção integral das instalações elétricas/mecânicas (118 artigos → ABNT/Bombeiros/"
       "Código de Edificações) e de matérias federais (segurança do trabalho, pesos e medidas).")
bullet("Fiscalização modernizada (notificação eletrônica, termo de compromisso) e penalidades "
       "em UFESP com dosimetria, embargo, interdição, lacração, emparedamento e cassação.")

# ============================================================ PARTE II
doc.add_heading("Parte II — Leitura estratégica sob o mandato", level=1)
p("Enquadramento: este é, na espinha, um projeto desregulatório — a nossa própria agenda. A "
  "jogada não é obstruir, é aprovar o eixo desburocratizante e cobrar caro pelos apêndices "
  "estatizantes (multa, poder coercitivo, delegação em branco). Enquadramento público: “a "
  "oposição pegou o código do governo e o melhorou para proteger o bolso do cidadão”.")
doc.add_heading("Apoiar e reivindicar", level=2)
bullet("Corte de ~48% dos artigos e fim de minúcias obsoletas; liberdade de horário (art. 229); "
       "saída de matérias de competência federal; tratamento diferenciado a MEI/ME/EPP (art. 300, "
       "V); termo de compromisso (arts. 293–298).")
doc.add_heading("Atacar ou condicionar", level=2)
bullet("Impacto fiscal sobre o contribuinte: multa diária de terreno sem teto (art. 101); taxa "
       "de administração de 20%→100%; multas de até 20.000 UFESP (drenagem).")
bullet("Estado coercitivo: interdição, lacração e emparedamento com recurso sem efeito "
       "suspensivo; demolição sem ordem judicial (art. 330, §2º); cassação de 3→5 anos "
       "alcançando sócios (art. 327, §2º).")
bullet("Delegação em branco ao Executivo (praias, calçadas, ambulantes, prazos).")
bullet("Novas obrigações ao setor privado (drenagem, câmeras, concessionárias).")
bullet("Viés principiológico do Estatuto da Cidade (arts. 2º–3º), a neutralizar com regra "
       "interpretativa e com o contrapeso da livre iniciativa (art. 3º, XI) e do art. 332.")

# ============================================================ PARTE III
doc.add_heading("Parte III — Análises temáticas aprofundadas", level=1)

doc.add_heading("III.1 Delegação a decreto: esvaziamento da função legislativa", level=2)
p("O PLC delega ao Executivo, sem parâmetros em lei: uso/barracas de praia (arts. 147, 149, "
  "150, 154); largura da faixa livre da calçada (art. 167, I) — que hoje está na lei; área e "
  "horário de cães na praia (art. 206); prazos de notificação (art. 288, VII); parklets (art. "
  "176); regras de parques (art. 157); valores de apreensão (art. 315).")
p("Opinião: nem toda delegação é vício — detalhe operacional pode e deve ir a regulamento. O "
  "problema é delegar a medida da restrição ao direito do cidadão, o que toca a reserva legal "
  "(CF, art. 5º, II e art. 37): restrição de direito e definição de sanção têm o núcleo "
  "essencial reservado à lei. Delegar isso em branco é deslegalização e transforma a Câmara em "
  "homologatória. Combatemos seletivamente: fixar em lei o quanto (pisos, critérios, prazos "
  "mínimos) e delegar apenas o como; e exigir as minutas de decreto antes da votação. "
  "(Emendas nº 5, 6 e 7.)")

doc.add_heading("III.2 Arquitetura das multas", level=2)
p("Problemas: faixas largas demais (alimentos e hospitais em 60–600 UFESP; drenagem em 20 a "
  "20.000 UFESP) geram discricionariedade e insegurança; valor fixo sem proporcionalidade "
  "(incêndio, 405 UFESP fixos, art. 203); multa diária sem teto (terreno, art. 101); risco de "
  "empilhamento (art. 302). Para ser justo, a dosimetria em si está na lei (atenuantes e "
  "agravantes, arts. 307–308).")
p("Como reduzir a quantidade e o peso das multas: escalonamento educativo/dupla visita para "
  "MEI/ME/EPP; descriminalizar condutas que não deveriam ser infração (Parte III.3); tabela "
  "única graduada por gravidade com multiplicadores objetivos; reparação no lugar de multa "
  "quando o dano for reparável; teto de acumulação da multa diária; vedação de bis in idem. "
  "(Emendas nº 8 a 11.)")

doc.add_heading("III.3 Interferência estatal sem necessidade", level=2)
bullet("Proibição de fabricar/comercializar fogos (art. 139, §2º): pune a atividade lícita, não "
       "a conduta lesiva (já vedada no caput), e provavelmente invade competência federal "
       "(produtos controlados pelo Exército — Decreto 10.030/2019 [Verificar]). Emenda nº 12.")
bullet("Proibição de animais em guarda/segurança (art. 209, II): banimento de atividade lícita; "
       "regular o bem-estar em vez de proibir. Emenda nº 13.")
bullet("Captura do corrupto (Callichirus) proibida em toda a orla, com apreensão (art. 153): "
       "regular período/quantidade seria proporcional. Emenda nº 14.")
bullet("Exposição de mercadoria fora do estabelecimento com apreensão (art. 230): atinge o "
       "pequeno comércio; advertência antes de apreender. Emenda nº 15.")
bullet("Câmeras obrigatórias (art. 212): terceiriza a vigilância ao privado, com custo e sem "
       "LGPD. Emenda nº 19.")
bullet("Charretes/carroças proibidas em todo o Município (art. 210, III): há trade-off de "
       "bem-estar animal — sinalizar, sem peitar frontalmente (escolher batalhas).")

doc.add_heading("III.4 Outras análises", level=2)
bullet("Fiscalização virando arrecadação: taxa de 100%, multa diária, leilão — exigir relatório "
       "anual público de multas aplicadas e arrecadadas (transparência é a nossa bandeira).")
bullet("Ingresso em domicílio: uniformizar zoonoses (arts. 59, 63) ao art. 283 e à CF, art. 5º, "
       "XI — só com ordem judicial, salvo flagrante ou risco. Emenda nº 17.")
bullet("Custo de conformidade ao pequeno negócio: empilhamento de exigências (drenagem, "
       "câmeras, caixa de gordura, CCT/ARTs) — dispensa/simplificação para pequeno porte. "
       "Emenda nº 10.")
bullet("Vacatio e transição: código entra em vigor na publicação (art. 335) — inseguro. "
       "Vacatio de 120 dias + período educativo + disposições transitórias. Emenda nº 18.")
bullet("Presunção de boa-fé e efeito suspensivo do recurso (arts. 319, 324): inverter o ônus a "
       "favor do particular. Emendas nº 16 e 20.")
bullet("Inflação normativa disfarçada: matérias jogadas para “legislação específica” futura — "
       "exigir que venham junto ou com prazo, para não legislarmos no escuro.")

# ============================================================ PARTE IV
doc.add_heading("Parte IV — Inovações não incorporadas pelo PLC", level=1)
p("O PLC não importa o instrumental pró-liberdade consolidado desde 2019. Quadro-síntese das "
  "inovações e da via de inserção (detalhe e fontes na Parte VII):")
tabela(
    ("#", "Inovação / emenda", "Base legal / exemplo"),
    (
        ("1", "Dispensa de licença/alvará p/ baixo risco + autodeclaração digital",
         "Lei 13.874/2019; CGSIM 51/2019 e 57/2020; SP, MG, BH"),
        ("2", "Silêncio positivo (deferimento tácito)", "Lei 13.874/2019, art. 3º"),
        ("3", "AIR + revisão quinquenal (sunset)", "Decreto 10.411/2020 (adaptar)"),
        ("4", "Once-only + alvará 100% digital", "Lei 14.129/2021 (Governo Digital)"),
        ("5", "Conformidade LGPD do videomonitoramento", "Lei 13.709/2018 (LGPD)"),
        ("6", "Regime de polo noturno + mediação + advertência antes da apreensão",
         "Proporcionalidade (LLE); experiências nacionais/internacionais"),
        ("7", "Micromobilidade (geofencing, estacionamento, responsabilidade)", "—"),
        ("8", "Poluição luminosa (diretriz) [avaliar custo]", "PL 1.400/2021; NBR 5101/2024"),
        ("9", "Requalificação de calçadas em vias estruturais [trade-off fiscal]",
         "SP Passeio Livre / PEC"),
        ("10", "Diretrizes pró-uso p/ parklets", "SP parklets"),
        ("11", "Dupla visita orientadora antes de multar MEI/ME/EPP", "Lei 13.874/2019, art. 55"),
    ),
    (0.7, 8.3, 7.0),
)
caixa("Estratégia: sua proposta de liberdade econômica já em tramitação [inserir nº do PL] vira "
      "a âncora da Emenda nº 1 — o Código passa a ser o veículo que a operacionaliza, sob a sua "
      "autoria (harmonizar, não duplicar).")

# ============================================================ PARTE V — EMENDAS
doc.add_page_break()
doc.add_heading("Parte V — Emendas (texto pronto para protocolo)", level=1)
p("Redação sugerida. Onde há “[ ]”, numerar conforme a posição final no texto do PLC. As "
  "emendas de reorganização estrutural (1 a 4) são candidatas a substitutivo/emenda coletiva "
  "da comissão; as demais podem ser individuais.", italic=True)

doc.add_heading("Bloco A — Liberdade econômica e governo digital", level=2)
emenda(1, "Aditiva",
       "Acréscimo de artigo ao Capítulo I do Título V (licenciamento), após o art. 221.",
       ["“Art. 221-A. Ficam dispensadas de licença de localização e funcionamento e de alvará "
        "prévios as atividades econômicas classificadas como de baixo risco, nos termos da Lei "
        "Complementar municipal nº [inserir nº do PL] e, subsidiariamente, da Resolução CGSIM "
        "nº 51/2019, bastando o registro e a autodeclaração eletrônica do interessado para o "
        "início e o pleno exercício da atividade.",
        "§ 1º A autodeclaração será emitida por meio digital, vedada a exigência de "
        "comparecimento presencial ou de documento já em poder da Administração.",
        "§ 2º É vedado exigir alvará, vistoria prévia ou taxa de licenciamento para as "
        "atividades de que trata o caput.”"],
       "Operacionaliza no Código a Lei federal 13.874/2019 e o PL de liberdade econômica em "
       "tramitação; corrige a omissão do art. 221, que mantém licença prévia para todos.")
emenda(2, "Aditiva",
       "Acréscimo de parágrafo ao art. 222.",
       ["“§ 7º Decorrido o prazo de 15 (quinze) dias sem decisão fundamentada do órgão "
        "competente, a licença considera-se deferida para todos os fins (aprovação tácita), "
        "ressalvadas as atividades de alto risco, hipótese em que o prazo poderá ser prorrogado "
        "uma única vez, por igual período, mediante despacho motivado.”"],
       "Institui o silêncio positivo (art. 3º da Lei 13.874/2019); hoje o §1º só permite operar "
       "com o protocolo, sem consequência à inércia do Executivo.")
emenda(3, "Aditiva",
       "Acréscimo de artigo às Disposições Finais.",
       ["“Art. [ ]. A criação, por decreto ou ato infralegal, de obrigação, exigência ou "
        "restrição a particular no âmbito deste Código dependerá de prévia Análise de Impacto "
        "Regulatório, com demonstração de necessidade, proporcionalidade e alternativas menos "
        "onerosas, nos moldes do Decreto federal nº 10.411/2020.",
        "Parágrafo único. As exigências regulamentares serão revistas a cada 5 (cinco) anos, "
        "ficando sem efeito as não expressamente confirmadas.”"],
       "Impõe crivo de custo/benefício às obrigações criadas por decreto e institui cláusula de "
       "caducidade (sunset) contra a inflação regulatória.")
emenda(4, "Modificativa e Aditiva",
       "Nova redação ao § 1º do art. 224 e acréscimo do § 3º.",
       ["“§ 1º O alvará será emitido em meio exclusivamente digital, dispensada a via física, e "
        "poderá ser apresentado à fiscalização por meio eletrônico ou código de verificação, "
        "vedada a exigência de afixação impressa.",
        "§ 3º É vedado exigir do interessado documento ou informação que já constem das bases de "
        "dados da Administração ou de outro órgão público, nos termos da Lei federal nº "
        "14.129/2021.”"],
       "Substitui o alvará afixado impresso por alvará digital e consagra o princípio "
       "“solicita uma única vez” (Governo Digital).")

doc.add_heading("Bloco B — Reserva legal e contenção da delegação", level=2)
emenda(5, "Modificativa",
       "Nova redação ao inciso I do art. 167.",
       ["“I - garantir uma faixa livre contínua, destinada exclusivamente à circulação de "
        "pedestres, desobstruída, com largura mínima de 1,50 m (um metro e cinquenta "
        "centímetros), podendo o regulamento apenas ampliá-la;”"],
       "Recoloca na lei o piso da faixa livre (hoje delegado a regulamento), preservando a "
       "acessibilidade e o controle legislativo.")
emenda(6, "Aditiva",
       "Acréscimo de parágrafo único ao art. 288.",
       ["“Parágrafo único. O prazo de que trata o inciso VII não será inferior a 15 (quinze) "
        "dias, salvo risco iminente à saúde ou à segurança, vedada sua redução por regulamento.”"],
       "Fixa piso legal ao prazo de defesa, impedindo que o decreto restrinja o direito ao "
       "contraditório.")
emenda(7, "Aditiva",
       "Acréscimo de artigo às Disposições Finais.",
       ["“Art. [ ]. Os regulamentos previstos neste Código limitar-se-ão a detalhar a execução "
        "de seus dispositivos, sendo-lhes vedado criar restrição, exigência ou sanção não "
        "prevista em lei.",
        "Parágrafo único. As minutas dos decretos regulamentadores serão encaminhadas à Câmara "
        "Municipal previamente à sua edição, para conhecimento.”"],
       "Barra o regulamento autônomo restritivo de direitos e dá transparência prévia à Câmara "
       "sobre o que se está delegando.")

doc.add_heading("Bloco C — Racionalização das multas", level=2)
emenda(8, "Aditiva",
       "Acréscimo de parágrafo ao art. 101.",
       ["“§ 4º A multa diária de que trata este artigo fica limitada a 30 (trinta) dias de "
        "incidência, após o que a Administração adotará as medidas de execução cabíveis, vedada "
        "a continuidade da cominação diária.”"],
       "Impede o efeito confiscatório da multa diária sem teto.")
emenda(9, "Modificativa",
       "Nova redação ao art. 203.",
       ["“Art. 203. As infrações ao disposto neste capítulo sujeitarão o infrator a multa de "
        "100 (cem) a 405 (quatrocentas e cinco) UFESP, graduada conforme a gravidade e o risco "
        "gerado.”"],
       "Substitui o valor fixo por faixa graduada, restaurando a proporcionalidade.")
emenda(10, "Aditiva",
       "Acréscimo de artigo ao Título VII (penalidades).",
       ["“Art. [ ]. Nas infrações de que resulte dano reparável, a autoridade fixará prazo para "
        "reparação, aplicando-se a multa apenas em caso de não reparação no prazo.",
        "Parágrafo único. Tratando-se de microempreendedor individual, microempresa ou empresa "
        "de pequeno porte, a primeira infração de natureza leve ensejará advertência orientadora "
        "(dupla visita), aplicando-se multa somente na reincidência, nos termos do art. 55 da "
        "Lei federal nº 13.874/2019.”"],
       "Prioriza a reparação sobre a arrecadação e protege o pequeno negócio, reduzindo o número "
       "de autuações.")
emenda(11, "Modificativa",
       "Nova redação ao art. 302.",
       ["“Art. 302. Quando de uma mesma conduta resultarem duas ou mais infrações, aplicar-se-á "
        "a penalidade da infração mais grave, vedada a cumulação de multas pelo mesmo fato.”"],
       "Veda o empilhamento de multas (non bis in idem) a partir de uma única conduta.")

doc.add_heading("Bloco D — Contenção da interferência estatal", level=2)
emenda(12, "Supressiva",
       "Supressão do § 2º do art. 139.",
       ["“Suprima-se o § 2º do art. 139, mantida a vedação de uso de fogos de estampido "
        "prevista no caput.”"],
       "Preserva a livre iniciativa: o mal a coibir é o uso ruidoso (já vedado), não a "
       "existência da fábrica/comércio, matéria ademais de competência federal (Exército).")
emenda(13, "Modificativa",
       "Nova redação ao inciso II do art. 209, com acréscimo do art. 209-A.",
       ["“II - (suprimido).”",
        "“Art. 209-A. Os estabelecimentos que utilizem animais em serviço de guarda ou "
        "vigilância observarão as normas de bem-estar animal e de registro definidas em "
        "regulamento, vedados os maus-tratos.”"],
       "Substitui o banimento de atividade lícita por regulação de bem-estar — regula, não "
       "proíbe.")
emenda(14, "Modificativa",
       "Nova redação ao art. 153.",
       ["“Art. 153. A captura dos crustáceos do gênero Callichirus será permitida mediante "
        "observância de período, quantidade e áreas definidos em regulamento, com vistas à "
        "sustentabilidade do estoque.”"],
       "Troca a proibição total por regulação proporcional, preservando atividade tradicional "
       "de subsistência.")
emenda(15, "Modificativa",
       "Nova redação ao § 2º do art. 230.",
       ["“§ 2º A infração sujeitará o responsável a advertência para regularização imediata e, "
        "na reincidência, a multa, reservando-se a apreensão às hipóteses de recusa persistente "
        "ou de obstrução da via pública.”"],
       "Afasta a apreensão como primeira medida contra o pequeno comércio, em favor da "
       "advertência.")

doc.add_heading("Bloco E — Garantias, LGPD e transição", level=2)
emenda(16, "Modificativa",
       "Nova redação ao § 3º do art. 319 e ao parágrafo único do art. 324.",
       ["“O recurso administrativo terá efeito suspensivo, salvo motivada situação de risco "
        "iminente à saúde, à segurança ou ao meio ambiente.”"],
       "Restabelece o efeito suspensivo como regra, exigindo do Estado a motivação do risco.")
emenda(17, "Aditiva",
       "Acréscimo de parágrafo ao art. 59 (controle de zoonoses).",
       ["“§ [ ]. O ingresso em imóvel residencial dependerá de consentimento do morador ou de "
        "ordem judicial, ressalvados o flagrante e o risco iminente, nos termos do art. 5º, XI, "
        "da Constituição Federal.”"],
       "Uniformiza o capítulo de zoonoses à garantia do domicílio e ao art. 283 do próprio PLC.")
emenda(18, "Modificativa e Aditiva",
       "Nova redação ao art. 335 e acréscimo de artigo de transição.",
       ["“Art. 335. Esta Lei Complementar entra em vigor 120 (cento e vinte) dias após a "
        "publicação, período durante o qual a fiscalização terá caráter exclusivamente "
        "orientador.”",
        "“Art. [ ]. Ficam preservadas as licenças e alvarás válidos até o respectivo termo; os "
        "processos e autos de infração em curso reger-se-ão pela norma vigente ao tempo do fato, "
        "aplicando-se retroatividade apenas quando mais benéfica ao administrado.”"],
       "Dá segurança jurídica à transição de um código de grande porte e institui período "
       "educativo.")
emenda(19, "Aditiva",
       "Acréscimo de parágrafo ao art. 212.",
       ["“§ [ ]. O tratamento das imagens observará a Lei federal nº 13.709/2018 (LGPD), com "
        "finalidade específica, controle de acesso, prazo de guarda e garantia dos direitos do "
        "titular, vedada a utilização para fim diverso.”"],
       "Fecha a lacuna de proteção de dados na videovigilância privada imposta pelo Código.")
emenda(20, "Aditiva",
       "Acréscimo de artigo ao Título I (disposições gerais).",
       ["“Art. [ ]. Presume-se a boa-fé do particular. As obrigações e restrições decorrentes "
        "deste Código somente serão exigíveis quando previstas expressamente em lei ou em "
        "regulamento a ela limitado, não gerando os princípios dos arts. 2º e 3º, por si sós, "
        "obrigação autônoma exigível do particular.”"],
       "Neutraliza a carga principiológica do Estatuto da Cidade e consagra a presunção de "
       "boa-fé (art. 2º da Lei 13.874/2019).")

# ============================================================ PARTE VI
doc.add_heading("Parte VI — Estratégia de veiculação e sequenciamento", level=1)
bullet("Substitutivo / emenda coletiva da comissão: emendas 1 a 7 (liberdade econômica, governo "
       "digital, reserva legal) — reorganizam o texto e ficam com autoria da comissão.")
bullet("Emendas individuais: 8 a 20 (multas, interferência estatal, garantias, LGPD, transição) "
       "— inserções pontuais, com a sua assinatura.")
bullet("Sequenciamento: priorizar 1–5, 10, 12, 18 e 19 — alto retorno, baixo custo político e "
       "coerentes com a bandeira do contribuinte; tratar 8–9, 11, 13–17 e 20 como bloco de "
       "aprofundamento; avaliar 6–9 da Parte IV (poluição luminosa e calçadas) só com crivo "
       "fiscal.")
bullet("Enquadramento público único: “o governo trocou a burocracia velha por um Estado mais "
       "discricionário, mais caro e ainda proibitivo — a oposição corrigiu as três coisas.”")

# ============================================================ PARTE VII
doc.add_heading("Parte VII — Diligências pendentes e fontes", level=1)
tabela(
    ("Diligência [Verificar]", "Por que importa"),
    (
        ("Nº do PL de liberdade econômica de sua autoria", "Citar na Emenda nº 1"),
        ("Valor da UFESP vigente", "Converter as multas em reais e dimensionar o impacto"),
        ("Norma federal exata dos fogos (Exército) e a distância de 200 m dos postos",
         "Fundamentar as Emendas nº 12 e a lacuna de inflamáveis"),
        ("Status das LCs alteradoras não listadas no art. 334",
         "Segurança jurídica; possível emenda à cláusula revogatória"),
    ),
    (7.0, 9.0),
)
doc.add_heading("Base normativa", level=2)
for t in (
    "Lei federal 13.874/2019 — Liberdade Econômica (baixo risco; aprovação tácita; dupla visita, "
    "art. 55; presunção de boa-fé).",
    "Resoluções CGSIM nº 51/2019 e 57/2020 — classificação de risco.",
    "Lei federal 14.129/2021 — Governo Digital (solicitação única; interoperabilidade).",
    "Decreto federal 10.411/2020 — Análise de Impacto Regulatório (referência a adaptar).",
    "Lei federal 13.709/2018 — LGPD; ABNT NBR 5101/2024 e PL federal 1.400/2021 — iluminação/"
    "poluição luminosa.",
    "Exemplos municipais: São Paulo (dispensa ~911 atividades; Passeio Livre/PEC; parklets), "
    "Belo Horizonte (275), Minas Gerais (915), Balneário Camboriú (1.114 de baixo risco).",
):
    bullet(t)
p("")
p("Documento de trabalho interno. Baseado no texto do PLC PA 59314/2025-78 e da Lei 3.531/1968, "
  "e nos entregáveis comparativo_posturas.xlsx e relatorio_posturas.docx. Itens de criação de "
  "obrigação/despesa devem passar pelo crivo fiscal do mandato antes de propostos. Redações de "
  "emenda sujeitas a ajuste técnico-legislativo final.", italic=True, size=9)

doc.save(SAIDA)
print(f"OK — {SAIDA.name}")
